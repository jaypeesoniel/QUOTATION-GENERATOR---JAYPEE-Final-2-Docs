import os
import tempfile
import shutil
from datetime import datetime
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

# ====== PATH NG IISA MONG QUOTATION FOLDER (optional local save) ======
GENERATED_FOLDER = os.path.join(os.getcwd(), "generated_docs")
# ==========================================================

# ✅ Gumamit ng relative path para gumana sa Render
TEMPLATES_FOLDER = os.path.join(os.getcwd(), "quotation_templates")

# --- Function to replace placeholders ---
def replace_placeholder(doc, placeholder, new_text, font_name="Cambria", font_size=12, bold=False):
    """Palitan ang placeholder sa buong document (kasama tables)."""
    for p in doc.paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.bold = bold
    # Tables support
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, new_text, font_name, font_size, bold)

@app.route('/', methods=['GET', 'POST'])
def index():
    # Auto-detect available templates
    if not os.path.exists(TEMPLATES_FOLDER):
        os.makedirs(TEMPLATES_FOLDER)

    models = [f.replace(".docx", "") for f in os.listdir(TEMPLATES_FOLDER) if f.endswith(".docx")]

    if request.method == 'POST':
        model = request.form.get('model_name', '').strip()
        client = request.form.get('client_name', '').strip().upper()
        address = request.form.get('address', '').strip().upper()
        contact = request.form.get('contact_person', '').strip().upper()

        # Current date
        current_date = datetime.now().strftime("%B %d, %Y")

        # Template check
        template_path = os.path.join(TEMPLATES_FOLDER, f"{model}.docx")
        if not os.path.exists(template_path):
            return f"❌ Walang template file para sa model '{model}'.<br><a href='/'>Balik</a>"

        # Format multi-line address
        address_lines = address.splitlines()
        formatted_address = "\n".join(line.strip() for line in address_lines if line.strip())

        # Temporary files
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "quotation.docx")
        output_filename = f"{model} - {client}.docx"
        temp_output = os.path.join(temp_dir, output_filename)

        # Load template and replace placeholders
        doc = Document(template_path)
        replace_placeholder(doc, "{{CLIENT_NAME}}", client, bold=True)
        replace_placeholder(doc, "{{ADDRESS}}", formatted_address)
        replace_placeholder(doc, "{{CONTACT_PERSON}}", contact, bold=True)
        replace_placeholder(doc, "{{DATE}}", current_date)
        replace_placeholder(doc, "ATTENTION:", "ATTENTION:", bold=True)
        doc.save(temp_docx)

        # Save to output file (no PDF conversion in Render)
        shutil.copy(temp_docx, temp_output)

        # ====== AUTO-SAVE SA IYONG SPECIFIC FOLDER ======
        os.makedirs(GENERATED_FOLDER, exist_ok=True)
        shutil.copy(temp_output, os.path.join(GENERATED_FOLDER, output_filename))
        # =================================================

        # Send file for browser download
        return send_file(temp_output, as_attachment=True, download_name=output_filename)

    return render_template('form.html', models=models)


if __name__ == '__main__':
    # Para gumana sa Render environment
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
